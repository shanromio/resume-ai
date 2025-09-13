from flask import Flask, render_template, request, redirect, url_for, session, send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import pdfplumber
import os
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
from datetime import datetime

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY")
app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv("DATABASE_URL")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

# ✅ Create tables inside app context
with app.app_context():
    db.create_all()

# Gemini setup
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# --- Models ---
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(80))
    email = db.Column(db.String(120), unique=True)
    password = db.Column(db.String(200))

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer)
    original_text = db.Column(db.Text)
    transformed_text = db.Column(db.Text)
    score = db.Column(db.Float)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

# --- Routes ---
@app.route("/")
def home():
    return redirect(url_for("signup"))

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        name = request.form["name"]
        email = request.form["email"]
        password = generate_password_hash(request.form["password"])
        user = User(name=name, email=email, password=password)
        db.session.add(user)
        db.session.commit()
        return redirect(url_for("login"))
    return render_template("signup.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]
        user = User.query.filter_by(email=email).first()
        if user and check_password_hash(user.password, password):
            session["user_id"] = user.id
            return redirect(url_for("index"))
        else:
            return render_template("login.html", error="Invalid credentials")
    return render_template("login.html")

@app.route("/index", methods=["GET", "POST"])
def index():
    if "user_id" not in session:
        return redirect(url_for("login"))

    if request.method == "POST":
        resume_file = request.files["resume"]
        job_desc = request.form["job_desc"]

        resume_text = ""
        with pdfplumber.open(resume_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    resume_text += text + "\n"

        # Keyword matching for ATS score
        job_keywords = [word.lower() for word in job_desc.split() if word.isalpha()]
        score = sum(1 for kw in job_keywords if kw in resume_text.lower())
        match_percent = round((score / len(job_keywords)) * 100, 2)

        # Gemini transformation
        prompt = f"""
        Rewrite the following resume to better match the job description below.
        Keep it professional, ATS-friendly, and tailored to the role.

        Resume:
        {resume_text}

        Job Description:
        {job_desc}
        """
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        transformed_resume = response.text.strip()

        # Save to database
        db.session.add(Resume(user_id=session["user_id"], original_text=resume_text,
                              transformed_text=transformed_resume, score=match_percent))
        db.session.commit()

        # Store for editing
        session["transformed_resume"] = transformed_resume
        session["score"] = match_percent

        return render_template("index.html", transformed=True,
                               transformed_resume=transformed_resume, score=match_percent)

    return render_template("index.html")

@app.route("/finalize", methods=["POST"])
def finalize():
    final_text = request.form["final_text"]

    # Save as .docx
    doc = Document()
    doc.add_heading("Final Resume", level=1)
    for line in final_text.split("\n"):
        doc.add_paragraph(line)
    doc.save("final_resume.docx")

    # Save for PDF
    session["final_text"] = final_text

    return redirect(url_for("download_options"))

@app.route("/download-options")
def download_options():
    return render_template("download_options.html")

@app.route("/download/docx")
def download_docx():
    return send_file("final_resume.docx", as_attachment=True)

@app.route("/download/pdf")
def download_pdf():
    final_text = session.get("final_text", "")
    pdf = FPDF()
    pdf.add_page()

    # ✅ Use Unicode-compatible font
    pdf.add_font("DejaVu", "", "fonts/DejaVuSans.ttf", uni=True)
    pdf.set_font("DejaVu", size=12)

    for line in final_text.split("\n"):
        pdf.multi_cell(0, 10, line)

    pdf.output("final_resume.pdf")
    return send_file("final_resume.pdf", as_attachment=True)


@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        return redirect(url_for("login"))
    resumes = Resume.query.filter_by(user_id=session["user_id"]).order_by(Resume.timestamp.desc()).all()
    return render_template("dashboard.html", resumes=resumes)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# --- Run App ---
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
    app.run(debug=True)
